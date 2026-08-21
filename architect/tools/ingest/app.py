"""Document ingestion worker — async queue (Redis) → OCR/chunk/embed → Qdrant.

Named collection: knowledge_chunks (never memory_N).
Supports workspace_id / thread_id for ACL-filtered RAG search.
"""
from __future__ import annotations

import asyncio
import hashlib
import json
import os
import re
import time
import uuid
from pathlib import Path
from typing import Any, Optional
from urllib.parse import unquote

import httpx
import redis
from fastapi import FastAPI, HTTPException, Query
from pydantic import BaseModel, Field

REDIS_URL = os.environ.get("REDIS_URL", "redis://redis:6379/0")
OCR_URL = os.environ.get("OCR_URL", "http://ocr:8091").rstrip("/")
EMBED_URL = os.environ.get("EMBED_URL", "http://embedding:8094").rstrip("/")
QDRANT_URL = os.environ.get("QDRANT_URL", "http://qdrant:6333").rstrip("/")
COLLECTION = os.environ.get("QDRANT_COLLECTION_KNOWLEDGE", "knowledge_chunks")
SECURITY_URL = os.environ.get("SECURITY_URL", "").rstrip("/")
AUTHZ_URL = os.environ.get("AUTHZ_URL", "http://authz:8097").rstrip("/")
NOTIFY_URL = os.environ.get("NOTIFY_URL", "http://notify:8092").rstrip("/")
# Bridge fallback when Notification Worker is off / unreachable (learn pending must still reach admin).
ZALO_BRIDGE_URL = (
    os.environ.get("ZALO_BRIDGE_URL")
    or os.environ.get("ZALO_PLUGIN_URL")
    or "http://zalo-proxy:8787"
).rstrip("/")
ZALO_PLUGIN_TOKEN = (os.environ.get("ZALO_PLUGIN_TOKEN") or "").strip()
ZALO_ADMIN_USERS_FILE = (
    os.environ.get("ZALO_ADMIN_USERS_FILE") or "/data/assistant/zalo_admin_users.txt"
).strip()
ZALO_ADMIN_USERS = (os.environ.get("ZALO_ADMIN_USERS") or "").strip()
NOTIFY_ZALO_THREAD = (os.environ.get("NOTIFY_ZALO_THREAD") or "").strip()
NOTIFY_ZALO_THREAD_TYPE = (
    os.environ.get("NOTIFY_ZALO_THREAD_TYPE") or "user"
).strip().lower() or "user"
QUEUE = "ingest:jobs"
PENDING_HASH = "ingest:pending"
# Shared lab job bus (also used by memory-manager index jobs)
MEMORY_QUEUE = os.environ.get("MEMORY_JOB_QUEUE", "memory:jobs")
MEDIA_ROOT = Path(os.environ.get("INGEST_MEDIA_ROOT", "/data/media"))
LEARN_DOCS_ROOT = Path(os.environ.get("LEARN_DOCS_ROOT", "/data/assistant/docs"))
LEARN_SCAN_MAX = int(os.environ.get("LEARN_SCAN_MAX", "200") or "200")
LEARN_REQUIRE_APPROVE = (os.environ.get("LEARN_REQUIRE_APPROVE") or "0").strip().lower() not in {
    "0",
    "false",
    "no",
    "off",
}
try:
    LEARN_LIST_LIMIT = int(os.environ.get("LEARN_LIST_LIMIT", "5") or "5")
except ValueError:
    LEARN_LIST_LIMIT = 5
# Admin-editable notify copy (JSON). Override path with LEARN_NOTIFY_PATH.
LEARN_NOTIFY_PATH = Path(
    os.environ.get("LEARN_NOTIFY_PATH", "/app/learn-notify.json")
)
ZALO_ALLOWED_USERS_FILE = Path(
    os.environ.get(
        "ZALO_ALLOWED_USERS_FILE",
        "/data/hermes/zalo_allowed_users.txt",
    )
)
LEARN_NOTIFY_DEFAULTS: dict[str, Any] = {
    "from_format": "{id}|{name}",
    "pending": {
        "title": "Knowledge — chờ duyệt",
        "body": (
            "id={id}\nfrom={from}\nname={name}\n{preview}\n"
            "Duyệt: !zalo learn approve {id}\n"
            "hoặc: !zalo learn approve {name}\n"
            "tất cả: !zalo learn approve *"
        ),
    },
    "approved": {"title": "Knowledge — đã học", "body": "{lines}"},
    "rejected": {"title": "Knowledge — từ chối", "body": "{lines}"},
    "deleted": {"title": "Knowledge — đã xóa", "body": "{lines}"},
    "scan": {"title": "Knowledge — scan docs", "body": "{lines}"},
}
LEARN_SCAN_EXTS = {
    ".pdf",
    ".xlsx",
    ".xlsm",
    ".xls",
    ".csv",
    ".tsv",
    ".docx",
    ".txt",
    ".md",
}

app = FastAPI(title="assistant-ingest", version="1.4.0")
r: redis.Redis | None = None
_worker: asyncio.Task | None = None


def _flow(stage: str, **fields: Any) -> None:
    parts = [f"[flow] stage={stage}"]
    for k, v in fields.items():
        if v is None:
            continue
        s = str(v).replace("\n", " ").replace('"', "'")
        if " " in s:
            s = f'"{s}"'
        parts.append(f"{k}={s}")
    print(" ".join(parts), flush=True)


class IngestReq(BaseModel):
    path: Optional[str] = None
    text: Optional[str] = None
    document_id: Optional[str] = None
    document_name: str = "document"
    source: str = "upload"
    workspace_id: Optional[str] = None
    thread_id: Optional[str] = None
    async_mode: bool = True


class ExtractTextReq(BaseModel):
    """Read office/CSV text now (agent turn), independent of the learn queue."""

    path: str
    max_chars: int = 0


class SearchReq(BaseModel):
    query: str
    workspace_id: Optional[str] = None
    thread_id: Optional[str] = None
    top_k: int = Field(default=6, ge=1, le=32)
    # If true, resolve thread→workspace via authz and filter
    acl: bool = True
    # Keep the highest-score chunk per document (catalog / cite)
    unique_document: bool = False


class LearnSubmit(BaseModel):
    text: Optional[str] = None
    path: Optional[str] = None
    document_name: str = "zalo-learn"
    source: str = "zalo"
    workspace_id: Optional[str] = None
    thread_id: Optional[str] = None
    sender_id: Optional[str] = None
    sender_name: Optional[str] = None


class LearnId(BaseModel):
    pending_id: str = ""
    selector: str = ""


class LearnScan(BaseModel):
    root: str = "docs"
    thread_id: Optional[str] = None
    sender_id: Optional[str] = None
    sender_name: Optional[str] = None


def _admin_zalo_dest() -> tuple[str, str]:
    """Resolve sole-admin Zalo DM (or NOTIFY_ZALO_THREAD override). Never log the uid."""
    tid = NOTIFY_ZALO_THREAD
    if tid:
        return tid, NOTIFY_ZALO_THREAD_TYPE if NOTIFY_ZALO_THREAD_TYPE in {"user", "group"} else "user"
    path = ZALO_ADMIN_USERS_FILE
    if path and Path(path).is_file():
        try:
            for line in Path(path).read_text(encoding="utf-8").splitlines():
                raw = line.strip()
                if not raw or raw.startswith("#"):
                    continue
                uid = raw.split("|", 1)[0].strip()
                if uid:
                    return uid, "user"
        except OSError:
            pass
    for part in ZALO_ADMIN_USERS.split(","):
        uid = part.strip()
        if uid:
            return uid, "user"
    return "", "user"


def _bridge_dm_admin(title: str, body: str) -> bool:
    """Direct bridge /send when notify worker is down — learn approve must not be silent."""
    tid, ttype = _admin_zalo_dest()
    if not tid or not ZALO_BRIDGE_URL:
        return False
    text = f"{title}\n{body}".strip()
    headers = {"content-type": "application/json"}
    if ZALO_PLUGIN_TOKEN:
        headers["x-bridge-token"] = ZALO_PLUGIN_TOKEN
        headers["Authorization"] = f"Bearer {ZALO_PLUGIN_TOKEN}"
    try:
        with httpx.Client(timeout=20) as c:
            r = c.post(
                f"{ZALO_BRIDGE_URL}/send",
                headers=headers,
                json={"threadId": tid, "threadType": ttype, "text": text[:3500]},
            )
            return r.status_code < 300
    except Exception:
        return False


def _notify_admin(title: str, body: str) -> bool:
    """Prefer Notification Worker; fall back to Zalo bridge DM so pending learn is never silent."""
    tid, ttype = _admin_zalo_dest()
    if NOTIFY_URL:
        try:
            payload: dict[str, Any] = {
                "title": title,
                "body": body,
                "severity": "info",
                "channels": ["zalo"],
                "kind": "alert",
            }
            if tid:
                payload["zalo_thread_id"] = tid
                payload["zalo_thread_type"] = ttype
            with httpx.Client(timeout=20) as c:
                r = c.post(f"{NOTIFY_URL}/v1/notify", json=payload)
                if r.status_code < 300:
                    data = r.json() if r.content else {}
                    if bool((data.get("results") or {}).get("zalo")):
                        return True
        except Exception:
            pass
    return _bridge_dm_admin(title, body)


def _learn_notify_cfg() -> dict[str, Any]:
    """Load editable templates; fall back to LEARN_NOTIFY_DEFAULTS."""
    cfg: dict[str, Any] = dict(LEARN_NOTIFY_DEFAULTS)
    path = LEARN_NOTIFY_PATH
    if not path.is_file():
        return cfg
    try:
        raw = json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return cfg
    if not isinstance(raw, dict):
        return cfg
    for key, val in raw.items():
        if key == "from_format" and isinstance(val, str) and val.strip():
            cfg["from_format"] = val.strip()
        elif isinstance(val, dict):
            base = dict(cfg.get(key) or {})
            if isinstance(val.get("title"), str) and val["title"].strip():
                base["title"] = val["title"].strip()
            if isinstance(val.get("body"), str) and val["body"].strip():
                base["body"] = val["body"]
            cfg[key] = base
    return cfg


def _format_template(tmpl: str, **fields: Any) -> str:
    out = tmpl or ""
    for key, val in fields.items():
        out = out.replace("{" + key + "}", str(val if val is not None else ""))
    return out


def _lookup_allowlist_name(sender_id: str) -> str:
    sid = (sender_id or "").strip()
    if not sid:
        return ""
    paths = (
        ZALO_ALLOWED_USERS_FILE,
        Path("/opt/data/zalo_allowed_users.txt"),
        Path("/data/hermes/zalo_allowed_users.txt"),
    )
    for path in paths:
        try:
            if not path.is_file():
                continue
            for line in path.read_text(encoding="utf-8", errors="replace").splitlines():
                raw = line.strip()
                if not raw or raw.startswith("#"):
                    continue
                uid, alias = (raw.split("|", 1) + [""])[:2]
                if uid.strip() == sid:
                    return alias.strip()
        except Exception:
            continue
    return ""


def _sender_from_label(sender_id: Optional[str], sender_name: Optional[str] = None) -> str:
    """Admin-facing from= line: id|displayName (UTF-8)."""
    sid = (sender_id or "").strip() or "unknown"
    name = (sender_name or "").strip()
    if not name or name == sid:
        name = _lookup_allowlist_name(sid)
    fmt = str((_learn_notify_cfg().get("from_format") or "{id}|{name}")).strip()
    if name and name != sid:
        return _format_template(fmt, id=sid, name=name)
    return sid


def _learn_notify(kind: str, **fields: Any) -> bool:
    cfg = _learn_notify_cfg()
    block = cfg.get(kind) if isinstance(cfg.get(kind), dict) else None
    if not isinstance(block, dict):
        block = LEARN_NOTIFY_DEFAULTS.get(kind) or {}
    title = str(block.get("title") or f"Knowledge — {kind}")
    body_tmpl = str(block.get("body") or "{lines}")
    body = _format_template(body_tmpl, **fields)
    return _notify_admin(title, body)


def _preview(text: str, n: int = 180) -> str:
    t = (text or "").replace("\n", " ").strip()
    if len(t) <= n:
        return t
    return t[: n - 1] + "…"


_GENERIC_TITLES = {
    "doc",
    "document",
    "file",
    "tài liệu",
    "tai lieu",
    "zalo-learn",
    "zalo learn",
}


def _human_doc_title(name: str) -> str:
    """Basename only: drop path, uuid prefix, extension. Not a server locator."""
    base = str(name or "").replace("\\", "/").split("/")[-1]
    base = re.sub(r"^(doc_)?[0-9a-f]{6,16}_", "", base, flags=re.I)
    base = re.sub(r"\.[A-Za-z0-9]{2,5}$", "", base)
    t = re.sub(r"[_\-]+", " ", base)
    t = re.sub(r"\s+", " ", t).strip(" .:-")
    if not t or t.lower() in _GENERIC_TITLES:
        return ""
    return t[:72]


def _looks_like_heading(s: str) -> bool:
    t = " ".join(str(s or "").split()).strip(" .:-•*")
    if len(t) < 12 or len(t) > 72:
        return False
    if not t[0].isalnum() or t[0].islower():
        return False
    if t.startswith("{") or re.match(r"^\d+(\.\d+)+\s", t):
        return False
    low = t.lower()
    if any(x in low for x in ("not found", "/opt/", "/data/", "inbound/")):
        return False
    words = t.split()
    if len(words) > 9:
        return False
    caps = sum(1 for w in words if w[:1].isupper())
    return caps >= max(1, len(words) // 3)


def _heading_from_preview(preview: str) -> str:
    raw = str(preview or "")
    for line in raw.replace("\r", "\n").split("\n")[:6]:
        line = re.sub(r"\s+", " ", line).strip(" .-•*")
        if _looks_like_heading(line):
            return line[:72]
    first = " ".join(raw.split())
    cut = first[:88]
    for sep in (". ", "? ", "! "):
        i = cut.find(sep)
        if 12 <= i <= 72:
            cut = cut[:i]
            break
    else:
        if len(first) > 72:
            cut = cut[:72].rsplit(" ", 1)[0] or cut[:72]
    cut = cut.strip(" .:-")
    return cut[:72] if _looks_like_heading(cut) else ""


def _public_title(name: str, preview: str = "") -> str:
    """Short label for chat: product/heading, never path or PDF body snippet."""
    human = _human_doc_title(name)
    if human:
        return human
    heading = _heading_from_preview(preview)
    if heading:
        return heading
    return "Tài liệu"


def _pending_get(pid: str) -> Optional[dict[str, Any]]:
    raw = _rdb().hget(PENDING_HASH, pid)
    if not raw:
        return None
    try:
        return json.loads(raw)
    except Exception:
        return None


def _pending_list() -> list[dict[str, Any]]:
    items: list[dict[str, Any]] = []
    for _k, raw in (_rdb().hgetall(PENDING_HASH) or {}).items():
        try:
            items.append(json.loads(raw))
        except Exception:
            continue
    items.sort(key=lambda x: float(x.get("submitted_at") or 0), reverse=True)
    return items


def _pending_put(item: dict[str, Any]) -> None:
    _rdb().hset(PENDING_HASH, item["pending_id"], json.dumps(item, ensure_ascii=False))


def _pending_del(pid: str) -> int:
    return int(_rdb().hdel(PENDING_HASH, pid))


def _match_pending(selector: str) -> list[dict[str, Any]]:
    """Resolve id | document name (exact then substring) | *."""
    sel = (selector or "").strip()
    items = _pending_list()
    if not sel:
        return []
    if sel in {"*", "all"}:
        return items
    exact_id = [it for it in items if str(it.get("pending_id") or "") == sel]
    if exact_id:
        return exact_id
    low = sel.lower()
    exact_name = [it for it in items if str(it.get("document_name") or "").lower() == low]
    if exact_name:
        return exact_name
    return [it for it in items if low in str(it.get("document_name") or "").lower()]


def _ingest_pending_item(item: dict[str, Any]) -> dict[str, Any]:
    pid = str(item.get("pending_id") or "")
    job = {
        "text": item.get("text") or "",
        "path": item.get("path"),
        "document_id": uuid.uuid4().hex,
        "document_name": item.get("document_name") or "zalo-learn",
        "rel_path": item.get("rel_path"),
        "source": item.get("source") or "zalo",
        "workspace_id": item.get("workspace_id"),
        "thread_id": item.get("thread_id"),
        "async_mode": False,
    }
    result = process_job(job)
    ok = bool(result.get("ok")) if isinstance(result, dict) else False
    chunks = result.get("chunks") if isinstance(result, dict) else None
    err = (result.get("error") if isinstance(result, dict) else None) or None
    if ok and pid:
        _pending_del(pid)
        _flow("learn_approved", pending_id=pid, document_id=job["document_id"], chunks=chunks)
    else:
        _flow("learn_approve_fail", pending_id=pid, error=err or "ingest failed")
    return {
        "pending_id": pid,
        "document_name": job["document_name"],
        "document_id": job["document_id"],
        "chunks": chunks,
        "ok": ok,
        "error": err,
        "ingest": result,
    }


def _rdb() -> redis.Redis:
    assert r is not None
    return r


def _resolve_media_path(path: str) -> Path:
    p = Path(path)
    if p.is_absolute():
        return p
    rel = path.lstrip("/")
    for base in (LEARN_DOCS_ROOT, MEDIA_ROOT):
        cand = base / rel
        if cand.is_file():
            return cand
    return MEDIA_ROOT / rel


SPREADSHEET_EXTS = {".xlsx", ".xlsm", ".xls"}
PLAIN_TEXT_EXTS = {".csv", ".tsv", ".md", ".txt", ".json", ".rst", ".html", ".xml", ".log", ".yaml", ".yml"}
DOC_EXTS = {".docx"}
SLIDE_EXTS = {".pptx"}
TEXT_EXTRACT_CHARS = 500000
SHEET_ROW_CAP = 2000


def _xlsx_text(p: Path) -> str:
    import openpyxl

    wb = openpyxl.load_workbook(p, read_only=True, data_only=True)
    parts: list[str] = []
    for ws in wb.worksheets:
        parts.append(f"## Sheet: {ws.title}")
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i >= SHEET_ROW_CAP:
                parts.append("...(truncated)")
                break
            parts.append("\t".join("" if c is None else str(c) for c in row))
    return "\n".join(parts)


def _docx_text(p: Path) -> str:
    import docx

    doc = docx.Document(str(p))
    parts = [para.text.strip() for para in doc.paragraphs if para.text and para.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append("\t".join(cells))
    return "\n".join(parts)


def _pptx_text(p: Path) -> str:
    """Slide text straight from the OOXML parts (no extra dependency)."""
    import re as _re
    import zipfile

    parts: list[str] = []
    with zipfile.ZipFile(p) as zf:
        slides = sorted(n for n in zf.namelist() if n.startswith("ppt/slides/slide") and n.endswith(".xml"))
        for i, name in enumerate(slides, start=1):
            xml = zf.read(name).decode("utf-8", "replace")
            runs = _re.findall(r"<a:t>(.*?)</a:t>", xml, flags=_re.S)
            lines = [_re.sub(r"\s+", " ", r).strip() for r in runs]
            body = "\n".join(line for line in lines if line)
            if body:
                parts.append(f"## Slide {i}\n{body}")
    return "\n\n".join(parts)


def _extract_text_from_path(path: str) -> str:
    """Best-effort text for office/CSV files before OCR."""
    p = _resolve_media_path(path)
    if not p.is_file():
        return ""
    ext = p.suffix.lower()
    try:
        if ext in SPREADSHEET_EXTS:
            return _xlsx_text(p)
        if ext in DOC_EXTS:
            return _docx_text(p)
        if ext in SLIDE_EXTS:
            return _pptx_text(p)
        if ext in PLAIN_TEXT_EXTS:
            return p.read_text(encoding="utf-8", errors="replace")[:TEXT_EXTRACT_CHARS]
    except Exception:
        return ""
    return ""


def _chunk(text: str, size: int = 1200, overlap: int = 150) -> list[str]:
    text = text.strip()
    if not text:
        return []
    out = []
    i = 0
    while i < len(text):
        out.append(text[i : i + size])
        i += max(1, size - overlap)
    return out


def _collection_vector_size(info: dict[str, Any]) -> Optional[int]:
    params = ((info.get("result") or {}).get("config") or {}).get("params") or {}
    vectors = params.get("vectors")
    if isinstance(vectors, dict) and isinstance(vectors.get("size"), int):
        return int(vectors["size"])
    return None


def _ensure_collection(dim: int = 1536) -> None:
    try:
        with httpx.Client(timeout=15) as c:
            resp = c.get(f"{QDRANT_URL}/collections/{COLLECTION}")
            if resp.status_code == 200:
                size = _collection_vector_size(resp.json() or {})
                if size == dim:
                    return
                if size is not None and size != dim:
                    c.delete(f"{QDRANT_URL}/collections/{COLLECTION}")
            c.put(
                f"{QDRANT_URL}/collections/{COLLECTION}",
                json={"vectors": {"size": dim, "distance": "Cosine"}},
            )
    except Exception:
        pass


def _embed(texts: list[str]) -> list[list[float]]:
    out: list[list[float]] = []
    batch_size = 16
    with httpx.Client(timeout=120) as c:
        for i in range(0, len(texts), batch_size):
            part = texts[i : i + batch_size]
            resp = c.post(f"{EMBED_URL}/v1/embeddings", json={"input": part})
            if resp.status_code >= 300:
                detail = ""
                try:
                    body = resp.json()
                    detail = str((body or {}).get("detail") or body)[:200]
                except Exception:
                    detail = (resp.text or "")[:200]
                raise RuntimeError(f"embed {resp.status_code}: {detail or resp.reason_phrase}")
            data = sorted((resp.json() or {}).get("data") or [], key=lambda x: x.get("index", 0))
            vecs = [d.get("embedding") for d in data if d.get("embedding")]
            if len(vecs) != len(part):
                raise RuntimeError(f"embed count {len(vecs)} != {len(part)}")
            out.extend(vecs)
    return out


def _upsert(points: list[dict]) -> None:
    with httpx.Client(timeout=60) as c:
        resp = c.put(f"{QDRANT_URL}/collections/{COLLECTION}/points", json={"points": points})
        resp.raise_for_status()


def _scroll_knowledge(limit_each: int = 128, max_points: int = 8000) -> list[dict[str, Any]]:
    out: list[dict[str, Any]] = []
    offset: Any = None
    with httpx.Client(timeout=30) as c:
        while len(out) < max_points:
            body: dict[str, Any] = {
                "limit": min(limit_each, max_points - len(out)),
                "with_payload": True,
                "with_vector": False,
            }
            if offset is not None:
                body["offset"] = offset
            resp = c.post(f"{QDRANT_URL}/collections/{COLLECTION}/points/scroll", json=body)
            if resp.status_code >= 300:
                break
            data = (resp.json() or {}).get("result") or {}
            batch = data.get("points") or []
            out.extend(batch)
            offset = data.get("next_page_offset")
            if not batch or offset is None:
                break
    return out


def _fold_learn_text(s: str) -> str:
    """ASCII-fold hyphens/spaces so citation copy-paste matches stored names."""
    t = s or ""
    try:
        t = unquote(t.replace("+", " "))
        if "%" in t:
            t = unquote(t)
    except Exception:
        pass
    t = re.sub(
        r"[\u00ad\u2010-\u2015\u2212\u2043\u058a\u1400\u1806\ufe58\ufe63\uff0d]",
        "-",
        t,
    )
    t = t.replace("\u00a0", " ").replace("\u202f", " ")
    return t


_LEARN_FILE_RE = re.compile(
    r"((?:[\w.\-]|%[0-9A-Fa-f]{2})+(?:[ \t]+(?:[\w.\-]|%[0-9A-Fa-f]{2})+)*"
    r"\.(?:pdf|docx?|xlsx?|xlsm|csv|txt|md|pptx?))",
    re.I,
)


def _learn_filenames(text: str) -> list[str]:
    folded = _fold_learn_text(text or "")
    out: list[str] = []
    seen: set[str] = set()
    for m in _LEARN_FILE_RE.finditer(folded):
        name = m.group(1).strip(".,;:()[]\u00ab\u00bb\"'")
        name = _fold_learn_text(name)
        if "." not in name or len(name) < 5:
            continue
        if name.lower() in {"skill.md", "readme.md"}:
            continue
        key = name.lower()
        if key not in seen:
            seen.add(key)
            out.append(name)
            stem = name.rsplit(".", 1)[0]
            for part in re.split(r"[-_\s]+", stem):
                pl = part.lower()
                if len(part) >= 8 and pl not in seen and pl not in {"skill", "readme", "document"}:
                    seen.add(pl)
                    out.append(part)
    return out


def _learn_excerpts(text: str) -> list[str]:
    """Quoted passages from a bot citation (content, not filenames)."""
    folded = _fold_learn_text(text or "")
    found: list[str] = []
    seen: set[str] = set()
    for rx in (
        r"[\u201c\u00ab\"]([^\u201d\u00bb\"]{20,400})[\u201d\u00bb\"]",
        r"^>\s*(.{20,400})$",
    ):
        for m in re.finditer(rx, folded, flags=re.M):
            s = re.sub(r"\s+", " ", m.group(1)).strip(" -")
            key = s.lower()
            if (
                len(s) >= 20
                and key not in seen
                and not re.search(r"trích dẫn|nội dung trích dẫn|tài liệu ·", s, flags=re.I)
            ):
                seen.add(key)
                found.append(s)
    return found


def _sentence_needles(folded: str) -> list[str]:
    parts = re.split(r"[\n·]|Section\s+\d+", folded)
    out: list[str] = []
    seen: set[str] = set()
    for p in parts:
        p = re.sub(r"\s+", " ", p).strip(" -")
        key = p.lower()
        if (
            28 <= len(p) <= 220
            and key not in seen
            and not re.search(r"trích dẫn|tài liệu ·|nội dung trích dẫn", p, flags=re.I)
        ):
            seen.add(key)
            out.append(p)
        if len(out) >= 12:
            break
    return out


def _cite_list_topic(text: str) -> str:
    """Keyword from a bot cite reply: Khớp «labsolution» — 5/6 file."""
    m = re.search(
        r"(?i)khớp\s*[\u00ab\u201c\"]([^\u00bb\u201d\"]{1,80})[\u00bb\u201d\"]",
        text or "",
    )
    return (m.group(1) or "").strip() if m else ""


def _cite_list_bullets(text: str) -> list[str]:
    out: list[str] = []
    for m in re.finditer(r"(?m)^[•]\s+(.+)$", text or ""):
        s = re.sub(r"\s+", " ", m.group(1)).strip()
        if not s or s.lower().startswith("còn "):
            continue
        if s.lower() in _GENERIC_TITLES:
            continue
        out.append(s[:120])
    return out


def _learn_selectors(sel: str) -> list[str]:
    """Keyword, filenames, and quoted passages from a citation / reply."""
    raw = (sel or "").strip()
    if not raw or raw in {"*", "all"}:
        return []
    folded = _fold_learn_text(raw).strip()
    keys: list[str] = []
    seen: set[str] = set()

    def add(x: str) -> None:
        x = _fold_learn_text(x).strip()
        k = x.lower()
        if x and k not in seen and k not in {"*", "all"}:
            seen.add(k)
            keys.append(x)

    topic = _cite_list_topic(raw)
    if topic:
        add(topic)
        return keys
    bullets = _cite_list_bullets(raw)
    if bullets and re.search(r"(?i)kiến thức đã học|còn \s*\d+\s*file", raw):
        for b in bullets:
            add(b)
        if keys:
            return keys

    for n in _learn_filenames(raw):
        add(n)
    for e in _learn_excerpts(raw):
        add(e)
    if not keys:
        if len(folded) <= 180:
            add(folded)
        else:
            for n in _sentence_needles(folded):
                add(n)
    elif len(folded) > 64:
        for n in _sentence_needles(folded):
            add(n)
    return keys


def _key_in_text(low: str, text_l: str) -> bool:
    if not low or not text_l:
        return False
    if low in text_l:
        return True
    if len(low) <= 90:
        return False
    words = low.split()
    if len(words) < 6:
        return False
    for i in range(0, len(words) - 5, 3):
        win = " ".join(words[i : i + 8])
        if len(win) >= 20 and win in text_l:
            return True
    return False


def _point_matches_keyword(payload: dict[str, Any], sel: str) -> bool:
    low = _fold_learn_text(sel).lower()
    if not low:
        return False
    doc_id = str(payload.get("document_id") or "")
    name = _fold_learn_text(str(payload.get("document_name") or ""))
    text = _fold_learn_text(str(payload.get("text") or ""))
    chunk_id = str(payload.get("chunk_id") or "")
    if doc_id.lower() == low or chunk_id.lower() == low:
        return True
    if len(low) >= 8 and doc_id.lower().startswith(low):
        return True
    name_l = name.lower()
    if name_l == low or low in name_l:
        return True
    if "." in name_l and len(name_l) >= 8 and name_l in low:
        return True
    title_l = _public_title(name, text).lower()
    if title_l and title_l not in _GENERIC_TITLES:
        if title_l == low or (len(low) >= 8 and (low in title_l or title_l in low)):
            return True
    return _key_in_text(low, text.lower())


def _pending_item_matches(item: dict[str, Any], sel: str) -> bool:
    if not sel:
        return True
    low = _fold_learn_text(sel).lower()
    pid = str(item.get("pending_id") or "")
    name = _fold_learn_text(str(item.get("document_name") or ""))
    text = _fold_learn_text(str(item.get("text") or ""))
    if pid == sel or pid.lower() == low:
        return True
    if len(low) >= 8 and pid.lower().startswith(low):
        return True
    name_l = name.lower()
    if name_l == low or low in name_l:
        return True
    if "." in name_l and len(name_l) >= 8 and name_l in low:
        return True
    title_l = _public_title(name, text).lower()
    if title_l and title_l not in _GENERIC_TITLES:
        if title_l == low or (len(low) >= 8 and (low in title_l or title_l in low)):
            return True
    return _key_in_text(low, text.lower())


def _catalog_limit(raw: Any = None) -> int:
    if raw is None:
        raw = LEARN_LIST_LIMIT
    try:
        n = int(raw)
    except (TypeError, ValueError):
        n = 5
    return max(1, min(50, n))


def _expand_hits_to_documents(points: list[dict[str, Any]], hits: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """If a chunk matches, drop the whole source document."""
    doc_ids = {
        str((h.get("payload") or {}).get("document_id") or "").strip()
        for h in hits
    }
    doc_ids.discard("")
    if not doc_ids:
        return hits
    out: list[dict[str, Any]] = []
    seen: set[Any] = set()
    for pt in points:
        did = str((pt.get("payload") or {}).get("document_id") or "").strip()
        pid = pt.get("id")
        if did in doc_ids and pid not in seen:
            seen.add(pid)
            out.append(pt)
    return out


def _knowledge_catalog(sel: str = "", with_hits: bool = False, limit: Any = None) -> dict[str, Any]:
    q = (sel or "").strip()
    if q in {"*", "all"}:
        q = ""
    q_sels = _learn_selectors(q) if q else []
    cap = _catalog_limit(limit)
    try:
        points = _scroll_knowledge()
    except Exception as exc:
        return {
            "ok": False,
            "error": f"qdrant scroll: {type(exc).__name__}: {exc}"[:180],
            "query": q,
            "documents": [],
            "pending": [],
            "hits": [],
        }
    docs: dict[str, dict[str, Any]] = {}
    hits: list[dict[str, Any]] = []
    for pt in points:
        payload = pt.get("payload") or {}
        if q:
            q_fold = _fold_learn_text(q).lower()
            name_l = _fold_learn_text(str(payload.get("document_name") or "")).lower()
            in_quote = bool(name_l and "." in name_l and len(name_l) >= 8 and name_l in q_fold)
            if not in_quote and name_l:
                for one in q_sels:
                    tok = _fold_learn_text(one).lower()
                    if len(tok) >= 8 and (tok in name_l or name_l in tok):
                        in_quote = True
                        break
            if not in_quote and not any(
                _point_matches_keyword(payload, one) for one in (q_sels or [q])
            ):
                continue
        did = str(payload.get("document_id") or "") or str(pt.get("id") or "")
        name = str(payload.get("document_name") or "doc")
        rec = docs.setdefault(
            did,
            {
                "document_id": did,
                "document_id_short": did[:8] if did else "",
                "document_name": name,
                "chunks": 0,
                "preview": "",
                "thread_id": payload.get("thread_id"),
                "source": payload.get("source"),
            },
        )
        rec["chunks"] += 1
        chunk_text = str(payload.get("text") or "")
        chunk_prev = _preview(chunk_text, 120)
        old = rec.get("preview") or ""
        old_bad = (
            not old
            or old.startswith("{")
            or "not found" in old.lower()
            or "/opt/" in old.lower()
            or "inbound/" in old.lower()
        )
        new_ok = (
            len(chunk_prev) >= 20
            and not chunk_prev.startswith("{")
            and "not found" not in chunk_prev.lower()
        )
        if old_bad and new_ok:
            rec["preview"] = chunk_prev
        elif not rec["preview"]:
            rec["preview"] = chunk_prev
        rec["title"] = _public_title(name, rec.get("preview") or "")
        if with_hits and q and len(hits) < cap:
            hits.append(
                {
                    "document_id": did,
                    "document_name": name,
                    "title": rec["title"],
                    "chunk_id": payload.get("chunk_id"),
                    "preview": _preview(str(payload.get("text") or ""), 160),
                }
            )
    pending_out: list[dict[str, Any]] = []
    for it in _pending_list():
        if q and not any(_pending_item_matches(it, one) for one in (q_sels or [q])):
            continue
        pending_out.append(
            {
                "pending_id": it.get("pending_id"),
                "document_name": it.get("document_name"),
                "title": _public_title(
                    str(it.get("document_name") or ""),
                    _preview(it.get("text") or ""),
                ),
                "sender_name": it.get("sender_name") or it.get("sender_id"),
                "thread_id": it.get("thread_id"),
                "preview": _preview(it.get("text") or ""),
                "submitted_at": it.get("submitted_at"),
            }
        )
    doc_list = sorted(docs.values(), key=lambda d: str(d.get("document_name") or "").lower())
    total = len(doc_list)
    pending_total = len(pending_out)
    return {
        "ok": True,
        "query": q,
        "count": total,
        "total": total,
        "limit": cap,
        "truncated": total > cap or pending_total > cap,
        "chunk_hits": sum(int(d.get("chunks") or 0) for d in doc_list),
        "documents": doc_list[:cap],
        "pending": pending_out[:cap],
        "hits": hits[:cap] if with_hits else [],
        "looked_up": [n for n in _learn_filenames(q) if "." in n][:cap] if q else [],
    }


def _delete_qdrant_ids(ids: list[Any]) -> int:
    if not ids:
        return 0
    with httpx.Client(timeout=60) as c:
        resp = c.post(
            f"{QDRANT_URL}/collections/{COLLECTION}/points/delete",
            json={"points": ids, "wait": True},
        )
        resp.raise_for_status()
    return len(ids)


def _resolve_workspace(thread_id: Optional[str], workspace_id: Optional[str]) -> Optional[str]:
    if workspace_id:
        return workspace_id
    if not thread_id or not AUTHZ_URL:
        return None
    try:
        with httpx.Client(timeout=10) as c:
            resp = c.post(
                f"{AUTHZ_URL}/v1/resolve",
                json={"platform": "zalo", "thread_id": thread_id},
            )
            if resp.status_code < 300:
                data = resp.json()
                return data.get("workspace_id") or (data.get("workspace") or {}).get("id")
    except Exception:
        pass
    return None


def process_job(job: dict[str, Any]) -> dict[str, Any]:
    text = job.get("text") or ""
    doc_id = job.get("document_id") or uuid.uuid4().hex
    name = job.get("document_name") or "document"
    ws = _resolve_workspace(job.get("thread_id"), job.get("workspace_id"))
    if not text and job.get("path"):
        text = _extract_text_from_path(job["path"])
        if text.strip():
            _flow(
                "ingest_parse_done",
                document_id=doc_id,
                path=job.get("path"),
                chars=len(text),
                method="spreadsheet",
            )
        if not text.strip():
            ocr_path = str(_resolve_media_path(str(job["path"])))
            _flow("ingest_ocr_start", document_id=doc_id, path=ocr_path, collection=COLLECTION)
            try:
                with httpx.Client(timeout=180) as c:
                    resp = c.post(f"{OCR_URL}/v1/ocr", json={"path": ocr_path})
                    if resp.status_code < 300:
                        text = resp.json().get("text") or ""
            except Exception:
                text = ""
            _flow(
                "ingest_ocr_done",
                document_id=doc_id,
                path=ocr_path,
                chars=len(text or ""),
                ok=bool(text and text.strip()),
            )
    if not text.strip():
        _flow("ingest_fail", document_id=doc_id, error="empty_text", collection=COLLECTION)
        return {"ok": False, "error": "empty_text", "document_id": doc_id}
    chunks = _chunk(text)
    if not chunks:
        _flow("ingest_fail", document_id=doc_id, error="empty_chunks", collection=COLLECTION)
        return {"ok": False, "error": "empty_chunks", "document_id": doc_id}
    try:
        vectors = _embed(chunks)
        if not vectors or not vectors[0]:
            raise RuntimeError("embed returned empty")
        _ensure_collection(dim=len(vectors[0]))
        points = []
        for i, (ch, vec) in enumerate(zip(chunks, vectors)):
            cid = hashlib.md5(f"{doc_id}:{i}".encode()).hexdigest()
            payload = {
                "document_id": doc_id,
                "document_name": name,
                "chunk_id": f"chunk_{i:04d}",
                "section": f"part-{i+1}",
                "source": job.get("source", "upload"),
                "text": ch,
            }
            rel = str(job.get("rel_path") or "").replace("\\", "/").strip()
            if rel:
                payload["rel_path"] = rel
            if job.get("path"):
                payload["path"] = str(job.get("path"))
            if ws:
                payload["workspace_id"] = ws
            if job.get("thread_id"):
                payload["thread_id"] = job["thread_id"]
            points.append(
                {
                    "id": int(cid[:15], 16),
                    "vector": vec,
                    "payload": payload,
                }
            )
        _upsert(points)
    except Exception as exc:
        err = f"{type(exc).__name__}: {exc}"[:180]
        _flow("ingest_fail", document_id=doc_id, error=err, collection=COLLECTION)
        return {"ok": False, "error": err, "document_id": doc_id}
    _flow(
        "ingest_rag_cached",
        document_id=doc_id,
        document_name=name,
        chunks=len(chunks),
        collection=COLLECTION,
        source=job.get("source", "upload"),
        workspace_id=ws,
        cache="qdrant",
    )
    return {
        "ok": True,
        "document_id": doc_id,
        "chunks": len(chunks),
        "collection": COLLECTION,
        "workspace_id": ws,
    }


async def _worker_loop() -> None:
    while True:
        try:
            item = _rdb().blpop(QUEUE, timeout=5)
            if not item:
                await asyncio.sleep(0.2)
                continue
            job = json.loads(item[1])
            await asyncio.to_thread(process_job, job)
        except Exception:
            await asyncio.sleep(1)


@app.on_event("startup")
async def startup() -> None:
    global r, _worker
    r = redis.Redis.from_url(REDIS_URL, decode_responses=True)
    _worker = asyncio.create_task(_worker_loop())


@app.get("/health")
def health() -> dict[str, Any]:
    try:
        _rdb().ping()
        q = _rdb().llen(QUEUE)
        pending = _rdb().hlen(PENDING_HASH)
    except Exception:
        return {"ok": False, "error": "redis_unavailable"}
    return {"ok": True, "queue": q, "pending": pending, "collection": COLLECTION, "memory_queue": MEMORY_QUEUE}


@app.post("/v1/extract-text")
def extract_text(req: ExtractTextReq) -> dict[str, Any]:
    """Synchronous office/CSV text so the agent can summarize in the same turn."""
    p = _resolve_media_path(req.path)
    text = _extract_text_from_path(req.path)
    limit = max(1000, min(int(req.max_chars or TEXT_EXTRACT_CHARS), TEXT_EXTRACT_CHARS))
    _flow(
        "extract_text",
        path=req.path,
        found=p.is_file(),
        chars=len(text),
        kind=p.suffix.lower(),
    )
    if not p.is_file():
        raise HTTPException(404, "file not found")
    return {"ok": bool(text), "path": str(p), "chars": len(text), "text": text[:limit]}


@app.post("/v1/ingest")
def ingest(req: IngestReq) -> dict[str, Any]:
    job = req.model_dump()
    job["document_id"] = req.document_id or uuid.uuid4().hex
    job["enqueued_at"] = time.time()
    if req.async_mode:
        _rdb().rpush(QUEUE, json.dumps(job))
        _flow(
            "ingest_enqueue",
            document_id=job["document_id"],
            document_name=req.document_name,
            source=req.source,
            workspace_id=req.workspace_id,
            thread_id=req.thread_id,
            queue=QUEUE,
            cache_target=COLLECTION,
        )
        return {"ok": True, "queued": True, "document_id": job["document_id"]}
    try:
        return process_job(job)
    except Exception:
        raise HTTPException(500, "ingest failed") from None


@app.post("/v1/search")
def search(req: SearchReq) -> dict[str, Any]:
    """ACL-filtered knowledge search over knowledge_chunks."""
    ws = _resolve_workspace(req.thread_id, req.workspace_id) if req.acl else req.workspace_id
    try:
        vectors = _embed([req.query])
    except Exception:
        raise HTTPException(502, "embed failed") from None
    fetch_k = req.top_k
    if req.unique_document:
        fetch_k = min(32, max(req.top_k * 4, req.top_k))
    body: dict[str, Any] = {
        "vector": vectors[0],
        "limit": fetch_k,
        "with_payload": True,
    }
    if req.acl and ws:
        body["filter"] = {
            "must": [{"key": "workspace_id", "match": {"value": ws}}]
        }
    elif req.acl and req.thread_id and not ws:
        # Unknown workspace → empty (default deny for protected RAG)
        return {"ok": True, "hits": [], "workspace_id": None, "denied": False, "empty_reason": "no_workspace"}
    try:
        with httpx.Client(timeout=30) as c:
            resp = c.post(f"{QDRANT_URL}/collections/{COLLECTION}/points/search", json=body)
            resp.raise_for_status()
            hits = resp.json().get("result") or []
    except Exception:
        raise HTTPException(502, "qdrant search failed") from None
    out = []
    for h in hits:
        payload = h.get("payload") or {}
        out.append(
            {
                "score": h.get("score"),
                "document_id": payload.get("document_id"),
                "document_name": payload.get("document_name"),
                "text": payload.get("text"),
                "workspace_id": payload.get("workspace_id"),
                "source": payload.get("source"),
            }
        )
    if req.unique_document:
        grouped: dict[str, dict[str, Any]] = {}
        for item in out:
            key = str(item.get("document_id") or item.get("document_name") or "")
            old = grouped.get(key)
            if not old or float(item.get("score") or 0) > float(old.get("score") or 0):
                grouped[key] = item
        out = sorted(grouped.values(), key=lambda x: -float(x.get("score") or 0))[: req.top_k]
    return {"ok": True, "hits": out, "workspace_id": ws, "collection": COLLECTION}


@app.post("/v1/learn/submit")
def learn_submit(req: LearnSubmit) -> dict[str, Any]:
    """Zalo knowledge-learn: stage pending until admin approves."""
    text = (req.text or "").strip()
    if not text and not req.path:
        raise HTTPException(400, "text or path required")
    pid = uuid.uuid4().hex[:8]
    item = {
        "pending_id": pid,
        "status": "pending",
        "text": text,
        "path": req.path,
        "document_name": req.document_name or "zalo-learn",
        "source": req.source or "zalo",
        "workspace_id": req.workspace_id,
        "thread_id": req.thread_id,
        "sender_id": req.sender_id,
        "sender_name": req.sender_name,
        "submitted_at": time.time(),
    }
    _pending_put(item)
    who = _sender_from_label(req.sender_id, req.sender_name)
    name = item["document_name"]
    notified = _learn_notify(
        "pending",
        id=pid,
        name=name,
        preview=_preview(text) or "(file/OCR)",
        **{"from": who},
    )
    _flow("learn_pending", pending_id=pid, sender_id=req.sender_id, notified=notified)
    return {"ok": True, "status": "pending", "pending_id": pid, "notified": notified}


@app.get("/v1/learn/pending")
def learn_pending() -> dict[str, Any]:
    items = _pending_list()
    out = []
    for it in items:
        out.append(
            {
                "pending_id": it.get("pending_id"),
                "document_name": it.get("document_name"),
                "sender_name": it.get("sender_name") or it.get("sender_id"),
                "thread_id": it.get("thread_id"),
                "preview": _preview(it.get("text") or ""),
                "submitted_at": it.get("submitted_at"),
            }
        )
    return {"ok": True, "count": len(out), "items": out}


@app.get("/v1/learn/list")
def learn_list(
    q: str = Query(default=""),
    limit: Optional[int] = Query(default=None),
) -> dict[str, Any]:
    """Indexed documents (+ pending). Optional keyword filter. Capped by LEARN_LIST_LIMIT."""
    return _knowledge_catalog(q, with_hits=False, limit=limit)


@app.post("/v1/learn/find")
def learn_find(req: LearnId, limit: Optional[int] = Query(default=None)) -> dict[str, Any]:
    """Find indexed chunks + pending by id / name / keyword (top N documents)."""
    sel = (req.selector or req.pending_id or "").strip()
    if not sel or sel in {"*", "all"}:
        raise HTTPException(400, "keyword required")
    data = _knowledge_catalog(sel, with_hits=True, limit=limit)
    _flow(
        "learn_find",
        selector=sel,
        documents=data.get("count"),
        chunks=data.get("chunk_hits"),
        pending=len(data.get("pending") or []),
    )
    return data


@app.post("/v1/learn/approve")
def learn_approve(req: LearnId) -> dict[str, Any]:
    sel = (req.selector or req.pending_id or "").strip()
    matched = _match_pending(sel)
    if not matched:
        raise HTTPException(404, "pending not found")
    items: list[dict[str, Any]] = []
    errors: list[str] = []
    for item in matched:
        pid = str(item.get("pending_id") or "")
        try:
            row = _ingest_pending_item(item)
            items.append(row)
            if not row.get("ok"):
                errors.append(f"{pid}:{row.get('error') or 'ingest failed'}")
        except Exception as exc:
            msg = f"{type(exc).__name__}: {exc}"[:180]
            errors.append(f"{pid}:{msg}")
            _flow("learn_approve_fail", pending_id=pid, error=msg)
    ok_items = [it for it in items if it.get("ok")]
    if not ok_items:
        return {
            "ok": False,
            "selector": sel,
            "count": 0,
            "items": items,
            "errors": errors,
            "error": "; ".join(errors) or "ingest failed",
        }
    lines = [f"selector={sel}", f"count={len(ok_items)}"]
    for it in ok_items:
        lines.append(f"id={it['pending_id']} name={it['document_name']} chunks={it.get('chunks')}")
    if errors:
        lines.append("errors=" + ",".join(errors))
    _learn_notify("approved", lines="\n".join(lines))
    first = ok_items[0]
    return {
        "ok": True,
        "selector": sel,
        "count": len(ok_items),
        "items": ok_items,
        "errors": errors,
        "pending_id": first["pending_id"],
        "ingest": first.get("ingest"),
    }


@app.post("/v1/learn/reject")
def learn_reject(req: LearnId) -> dict[str, Any]:
    sel = (req.selector or req.pending_id or "").strip()
    matched = _match_pending(sel)
    if not matched:
        raise HTTPException(404, "pending not found")
    items: list[dict[str, Any]] = []
    for item in matched:
        pid = str(item.get("pending_id") or "")
        name = item.get("document_name")
        if pid:
            _pending_del(pid)
        _flow("learn_rejected", pending_id=pid, name=name)
        items.append(
            {
                "pending_id": pid,
                "document_name": name,
                "status": "rejected",
                "thread_id": item.get("thread_id") or "",
                "sender_id": item.get("sender_id") or "",
            }
        )
    lines = [f"selector={sel}", f"count={len(items)}"]
    for it in items:
        lines.append(f"id={it['pending_id']} name={it.get('document_name')}")
    # Never announce reject to the submitter. Admin ack is the !zalo reply
    # (and even that is suppressed when typed in the requester's chat).
    flag = (os.environ.get("LEARN_NOTIFY_REJECT") or "0").strip().lower()
    if flag in {"1", "true", "yes", "on"}:
        _learn_notify("rejected", lines="\n".join(lines))
    first = items[0]
    return {
        "ok": True,
        "selector": sel,
        "count": len(items),
        "items": items,
        "pending_id": first["pending_id"],
        "status": "rejected",
    }


@app.post("/v1/learn/delete")
def learn_delete(req: LearnId) -> dict[str, Any]:
    """Remove indexed chunks + matching pending by id / name / keyword. No * wipe."""
    sel = (req.selector or req.pending_id or "").strip()
    sels = _learn_selectors(sel)
    if not sels or any(s in {"*", "all"} for s in sels):
        raise HTTPException(400, "keyword required (not *)")
    sel = ", ".join(sels)
    if len(sel) > 160:
        sel = sel[:157] + "..."
    pending_removed: list[dict[str, Any]] = []
    seen_pending: set[str] = set()
    for one in sels:
        for item in _match_pending(one):
            pid = str(item.get("pending_id") or "")
            if pid in seen_pending:
                continue
            seen_pending.add(pid)
            name = item.get("document_name")
            if pid:
                _pending_del(pid)
            pending_removed.append({"pending_id": pid, "document_name": name, "status": "deleted"})
    try:
        points = _scroll_knowledge()
        hits = []
        seen_pt: set[Any] = set()
        for one in sels:
            for pt in points:
                pid = pt.get("id")
                if pid in seen_pt:
                    continue
                if _point_matches_keyword(pt.get("payload") or {}, one):
                    seen_pt.add(pid)
                    hits.append(pt)
        hits = _expand_hits_to_documents(points, hits)
    except Exception as exc:
        return {
            "ok": False,
            "selector": sel,
            "error": f"qdrant scroll: {type(exc).__name__}: {exc}"[:180],
            "pending": pending_removed,
        }
    ids = [pt.get("id") for pt in hits if pt.get("id") is not None]
    try:
        deleted = _delete_qdrant_ids(ids)
    except Exception as exc:
        return {
            "ok": False,
            "selector": sel,
            "error": f"qdrant delete: {type(exc).__name__}: {exc}"[:180],
            "pending": pending_removed,
            "points": 0,
        }
    docs = sorted(
        {
            str((pt.get("payload") or {}).get("document_name") or "").strip()
            for pt in hits
        }
        - {""}
    )
    if deleted == 0 and not pending_removed:
        return {
            "ok": True,
            "selector": sel,
            "count": 0,
            "points": 0,
            "documents": [],
            "pending": [],
            "empty": True,
        }
    lines = [f"selector={sel}", f"points={deleted}", f"pending={len(pending_removed)}"]
    if docs:
        lines.append("docs=" + ", ".join(docs[:12]))
    _learn_notify("deleted", lines="\n".join(lines))
    _flow("learn_deleted", selector=sel, points=deleted, pending=len(pending_removed), docs=",".join(docs[:8]) or None)
    return {
        "ok": True,
        "selector": sel,
        "count": deleted + len(pending_removed),
        "points": deleted,
        "documents": docs,
        "pending": pending_removed,
    }


_SCAN_SKIP_DIRS = {".git", ".trash", ".ds_store", "__pycache__", "node_modules"}


def _indexed_document_keys() -> set[str]:
    """Paths and names already in knowledge_chunks (lowercase)."""
    keys: set[str] = set()
    try:
        for pt in _scroll_knowledge():
            payload = pt.get("payload") or {}
            for field in ("rel_path", "document_name", "path"):
                val = str(payload.get(field) or "").strip().replace("\\", "/").lower()
                if val:
                    keys.add(val)
    except Exception:
        pass
    return keys


def _iter_docs_files(root: Path) -> list[Path]:
    out: list[Path] = []
    root = root.resolve()
    if not root.is_dir():
        return out
    for dirpath, dirnames, filenames in os.walk(root):
        dirnames[:] = [d for d in dirnames if d.lower() not in _SCAN_SKIP_DIRS and not d.startswith(".")]
        for name in filenames:
            if name.startswith(".") or name.startswith("~$"):
                continue
            p = Path(dirpath) / name
            if p.suffix.lower() not in LEARN_SCAN_EXTS:
                continue
            try:
                p.resolve().relative_to(root)
            except ValueError:
                continue
            if p.is_file():
                out.append(p)
    out.sort(key=lambda x: str(x).lower())
    return out


@app.post("/v1/learn/scan")
def learn_scan(req: LearnScan) -> dict[str, Any]:
    """Walk CloudDrive/docs mirror → pending learn. Skip already pending or indexed. No * wipe."""
    alias = (req.root or "docs").strip().lower() or "docs"
    if alias in {"doc"}:
        alias = "docs"
    if alias not in {"docs", "clouddrive"}:
        raise HTTPException(400, "root must be docs or clouddrive")
    if alias == "clouddrive":
        root = Path(os.environ.get("CLOUDDRIVE_MIRROR_DIR") or "/data/clouddrive")
    else:
        root = LEARN_DOCS_ROOT
    if not root.is_dir():
        _flow("learn_scan_fail", root=str(root), error="missing")
        return {
            "ok": False,
            "error": f"docs root missing ({root}) — mount CLOUDDRIVE_MIRROR_DIR / enable clouddrive sync",
            "root": str(root),
            "submitted": [],
            "skipped": [],
        }
    files = _iter_docs_files(root)
    if not files:
        _flow("learn_scan", root=str(root), files=0)
        return {
            "ok": True,
            "root": str(root),
            "empty": True,
            "scanned": 0,
            "count": 0,
            "submitted": [],
            "skipped": [],
        }
    pending_now = _pending_list()
    pending_paths = {str(it.get("path") or "").replace("\\", "/").lower() for it in pending_now}
    pending_rels = {str(it.get("rel_path") or "").replace("\\", "/").lower() for it in pending_now}
    indexed = _indexed_document_keys()
    submitted: list[dict[str, Any]] = []
    skipped: list[dict[str, str]] = []
    for f in files:
        if len(submitted) >= LEARN_SCAN_MAX:
            skipped.append({"name": f.name, "reason": "cap"})
            continue
        rel = f.resolve().relative_to(root.resolve()).as_posix()
        abs_path = str(f.resolve())
        name = rel  # unique per file; many skills share SKILL.md basename
        low_rel = rel.lower()
        low_abs = abs_path.replace("\\", "/").lower()
        if low_abs in pending_paths or low_rel in pending_rels:
            skipped.append({"name": name, "reason": "pending"})
            continue
        if low_rel in indexed or low_abs in indexed:
            skipped.append({"name": name, "reason": "indexed"})
            continue
        pid = uuid.uuid4().hex[:8]
        item = {
            "pending_id": pid,
            "status": "pending",
            "text": "",
            "path": abs_path,
            "document_name": name,
            "source": "docs",
            "workspace_id": None,
            "thread_id": req.thread_id,
            "sender_id": req.sender_id,
            "sender_name": req.sender_name,
            "submitted_at": time.time(),
            "rel_path": rel,
        }
        _pending_put(item)
        pending_paths.add(abs_path.replace("\\", "/").lower())
        pending_rels.add(low_rel)
        if not LEARN_REQUIRE_APPROVE:
            # Product default: auto-learn with no admin approve
            row = _ingest_pending_item(item)
            submitted.append(
                {
                    "pending_id": pid,
                    "document_name": name,
                    "rel_path": rel,
                    "auto": True,
                    "ok": bool(row.get("ok")),
                    "document_id": row.get("document_id"),
                }
            )
        else:
            submitted.append({"pending_id": pid, "document_name": name, "rel_path": rel})
    n = len(submitted)
    skip_n = len(skipped)
    lines = [
        f"root={root}",
        f"scanned={len(files)}",
        f"new={n}",
        f"skip={skip_n}",
        f"require_approve={int(LEARN_REQUIRE_APPROVE)}",
    ]
    for it in submitted[:12]:
        lines.append(f"id={it['pending_id']} {it['document_name']}")
    if n and LEARN_REQUIRE_APPROVE:
        lines.append("Duyệt: !zalo learn approve *")
    notified = False
    if n and LEARN_REQUIRE_APPROVE:
        notified = _learn_notify("scan", lines="\n".join(lines))
    _flow("learn_scan", root=str(root), scanned=len(files), new=n, skip=skip_n, notified=notified)
    return {
        "ok": True,
        "root": str(root),
        "scanned": len(files),
        "count": n,
        "submitted": submitted,
        "skipped": skipped,
        "notified": notified,
        "auto_ingest": not LEARN_REQUIRE_APPROVE,
        "capped": any(s.get("reason") == "cap" for s in skipped),
    }

