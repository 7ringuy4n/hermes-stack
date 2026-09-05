#!/usr/bin/env python3
"""Structural gate for styled PDF/Word/PowerPoint/Excel renderer inputs."""
from __future__ import annotations

import sys
import tempfile
from pathlib import Path

ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(ROOT / "architect/models/dispatcher"))

from office_file import write_docx_styled, write_pptx_styled, write_xlsx_styled  # noqa: E402

BODY = """# Quarterly overview
## A concise operational snapshot
- Availability: 99.9%
- Requests: 1,240

Context for the reporting period.

**Owner:** Operations

| Work item | Owner | Status |
|---|---|---|
| Review | Team lead | Ready |

### Highlights
- Stable delivery
- Clear ownership
"""


def main() -> int:
    from docx import Document
    from openpyxl import load_workbook
    from pptx import Presentation

    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        docx_path = write_docx_styled(root / "report.docx", BODY)
        xlsx_path = write_xlsx_styled(root / "report.xlsx", BODY)
        pptx_path = write_pptx_styled(root / "report.pptx", BODY)

        doc = Document(docx_path)
        wb = load_workbook(xlsx_path)
        ws = wb["Overview"]
        deck = Presentation(pptx_path)
        checks = {
            "docx has styled title": any(p.style.name == "Title" and p.text for p in doc.paragraphs),
            "docx has fact table": len(doc.tables) >= 2,
            "docx renders authored table": any(len(table.columns) == 3 for table in doc.tables),
            "docx strips markdown chrome": not any(
                marker in paragraph.text
                for marker in ("**", "|---|", "| Work item |")
                for paragraph in doc.paragraphs
            ),
            "xlsx title is merged": "A1:D1" in {str(rng) for rng in ws.merged_cells.ranges},
            "xlsx grid hidden": ws.sheet_view.showGridLines is False,
            "xlsx useful widths": ws.column_dimensions["A"].width >= 20,
            "xlsx renders authored table": ws.max_column >= 4 and any(
                cell.value == "Work item" for row in ws.iter_rows() for cell in row
            ),
            "pptx has multiple slides": len(deck.slides) >= 2,
            "pptx uses widescreen": deck.slide_width > deck.slide_height,
            "formats preserved": all(path.suffix in {".docx", ".xlsx", ".pptx"} for path in (docx_path, xlsx_path, pptx_path)),
            "dispatcher ships office renderers": all(
                package in (ROOT / "architect/models/dispatcher/Dockerfile").read_text(encoding="utf-8")
                for package in ("libreoffice-writer", "libreoffice-calc", "libreoffice-impress")
            ),
        }
    for name, ok in checks.items():
        print(("PASS" if ok else "FAIL"), name)
    return 0 if all(checks.values()) else 1


if __name__ == "__main__":
    raise SystemExit(main())
